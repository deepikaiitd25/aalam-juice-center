from docx import Document


def create_docx(title, sections, output_file="output.docx"):
    doc = Document()

    # Title
    doc.add_heading(title, 0)

    # Sections
    for section in sections:
        doc.add_heading(section["heading"], level=1)
        doc.add_paragraph(section["content"])

    doc.save(output_file)
    return output_file