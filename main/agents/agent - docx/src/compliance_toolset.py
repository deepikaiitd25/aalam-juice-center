from docx import Document

class DocumentToolset:

    def generate_docx(self, title, content, output_file="output.docx"):
        doc = Document()

        # Title
        doc.add_heading(title, 0)

        # Content
        for line in content.split("\n"):
            doc.add_paragraph(line)

        doc.save(output_file)
        return f"DOCX created: {output_file}"
