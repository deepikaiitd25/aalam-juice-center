import os
import json
import logging
import io
import re
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# --- THE MAGIC PARSER ---


def _add_rich_text(paragraph, text):
    """Parses **bold** and *italics* markdown and applies native Word formatting."""
    # Split by bold first, then italics
    bold_parts = re.split(r'(\*\*.*?\*\*)', str(text))
    for b_part in bold_parts:
        if b_part.startswith('**') and b_part.endswith('**'):
            # It's bold
            clean_text = b_part[2:-2]
            run = paragraph.add_run(clean_text)
            run.bold = True
        else:
            # Check for italics inside the non-bold parts
            italic_parts = re.split(r'(\*.*?\*)', b_part)
            for i_part in italic_parts:
                if i_part.startswith('*') and i_part.endswith('*'):
                    clean_text = i_part[1:-1]
                    run = paragraph.add_run(clean_text)
                    run.italic = True
                else:
                    paragraph.add_run(i_part)


class DocxToolset:
    def __init__(self, host: str, port: int):
        self.host = host
        self.port = port
        self.output_dir = "outputs"
        os.makedirs(self.output_dir, exist_ok=True)
        logger.info(
            f"Initialized DocxToolset. Saving files to ./{self.output_dir}")

    async def generate_docx(self, filename: str, title: str, sections: list) -> str:
        """Generate a highly formatted Word document with rich text, tables, and charts."""
        try:
            if isinstance(sections, str):
                sections = json.loads(sections)

            if not filename.endswith(".docx"):
                filename += ".docx"

            doc = Document()

            # --- DEFAULT STYLING ---
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)

            # --- TITLE PAGE ---
            doc.add_paragraph()
            doc.add_paragraph()
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()

            # --- CONTENT RENDERER ---
            for section in sections:
                if not isinstance(section, dict):
                    continue

                heading = section.get("heading", "")
                content = section.get("content", "")
                sec_type = section.get("type", "paragraph")

                # Render Heading
                if heading:
                    h = doc.add_heading(heading, level=1)
                    h.runs[0].font.color.rgb = RGBColor(
                        0x1F, 0x49, 0x7D)  # Professional Blue

                # Render Content based on type
                if sec_type == "list" and isinstance(content, list):
                    for item in content:
                        p = doc.add_paragraph(style='List Bullet')
                        # Apply formatting to list items
                        _add_rich_text(p, item)

                elif sec_type == "table" and isinstance(content, list):
                    rows = len(content)
                    cols = len(content[0]) if rows > 0 else 1
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Light Shading Accent 1'

                    for i, row_data in enumerate(content):
                        row_cells = table.rows[i].cells
                        for j, cell_data in enumerate(row_data):
                            p = row_cells[j].paragraphs[0]
                            _add_rich_text(p, cell_data)
                            # THE FIX: Auto-bold the header row!
                            if i == 0:
                                for run in p.runs:
                                    run.bold = True

                elif sec_type == "chart" and isinstance(content, dict):
                    labels = content.get("labels", [])
                    values = content.get("values", [])
                    chart_title = content.get("title", heading)

                    if labels and values:
                        plt.figure(figsize=(6, 4))
                        plt.bar(labels, values, color='#1F497D')
                        plt.title(chart_title)
                        plt.ylabel(content.get("y_label", ""))
                        plt.xlabel(content.get("x_label", ""))

                        memfile = io.BytesIO()
                        plt.savefig(memfile, format='png', bbox_inches='tight')
                        memfile.seek(0)
                        plt.close()

                        doc.add_picture(memfile, width=Inches(5.5))

                else:
                    # Standard paragraph with rich text and clean justification
                    p = doc.add_paragraph()
                    _add_rich_text(p, content)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)

            url = f"http://{self.host}:{self.port}/outputs/{filename}"
            return f"✅ Successfully generated **{filename}**!\n\n📥 [Download your professional report here]({url})"

        except Exception as e:
            logger.error(f"❌ DOCX Generation failed: {e}")
            return f"❌ Failed to generate report: {str(e)}"

    def get_tools(self):
        return {"generate_docx": self}
